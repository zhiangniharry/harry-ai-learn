# -*- coding: utf-8 -*-
"""生成 The Melting Point 原文版 + 中文版 docx"""
from docx import Document
from docx.shared import Pt, RGBColor

def style(doc):
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    try:
        doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/mathControls}mx', 'Microsoft YaHei')
    except:
        pass

def h1(doc, t):
    p = doc.add_heading(level=1)
    r = p.add_run(t)
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0,0,128)
    return p

def h2(doc, t):
    p = doc.add_heading(level=2)
    r = p.add_run(t)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0,80,0)
    return p

def para(doc, t, sz=11):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(sz)
    return p

# ===================== 1. 原文版 =====================
doc_en = Document()
style(doc_en)
h1(doc_en, "The Melting Point")
para(doc_en, "By Frank Gao - Dimension Research")
para(doc_en, "Source: https://research.dimensioncap.com/p/the-melting-point")
para(doc_en, "Published: 2026-03-23")
para(doc_en, "")

h2(doc_en, "Abstract")
para(doc_en, "We applied autoresearch to protein thermostability, giving AI agents two levels of freedom pitted against Bayesian optimization, and learned when intuition can synergize with exhaustion.")

h2(doc_en, "Background: Why Thermostability Matters")
para(doc_en, "Every protein has a melting temperature (Tm). Heat it past that threshold and the structure unravels. A therapeutic antibody that denatures below 55C cannot be formulated into a stable drug product. Thermostability correlates tightly with drug developability properties: solubility, aggregation resistance, shelf life, and manufacturability.")

para(doc_en, "Thermostability prediction has received comparatively less attention than binding affinity prediction. From a reductionist lens, Tm can be reliably measured and the readout is a positive scalar - making it a compelling first project for autoresearch in biology.")

h2(doc_en, "Dataset: Meltome Atlas + TemBERTureDB")
para(doc_en, "- Meltome Atlas: melting curves for 13,000+ proteins across 13 species")
para(doc_en, "- TemBERTureDB: cleaned and curated, balanced thermophilic/non-thermophilic sequences")
para(doc_en, "- Train/val/test split: 80/10/10, clustered to prevent leakage")
para(doc_en, "- Task: predict Tm from amino acid sequence, evaluate on MAE (degrees Celsius)")

h2(doc_en, "Baseline Model")
para(doc_en, "Reference: TemBERTure_Tm with MAE 6.42C using protBERT-BFD backbone + Houlsby bottleneck adapters (25M trainable params on 420M frozen backbone).")
para(doc_en, "Our model: frozen ESM-2 150M embeddings + lightweight self-attention adapter + MLP regression head (~1.7M params, 15x smaller).")

h2(doc_en, "Five Optimization Approaches")
para(doc_en, "1. Bayesian Optimization (BO): non-agentic baseline, searches ~25 hyperparameters via statistical model")
para(doc_en, "2. Restricted Agent (RA): same 25 knobs as BO, but LLM-controlled, proposes 1-3 changes per trial, tracks history")
para(doc_en, "3. Unleashed Agent (UA): no constraints, can rewrite entire training script, invent new architectures")
para(doc_en, "4. RA + BO: agent sets direction, BO exploits continuous subspace")
para(doc_en, "5. UA + BO: same synthesis applied to UA open-ended proposals")

h2(doc_en, "Key Findings")
para(doc_en, "1. All methods beat baseline with 15x smaller model")
para(doc_en, "2. UA achieved lowest validation MAE; RA+BO achieved lowest test MAE (best generalization)")
para(doc_en, "3. Paradox of freedom: UA tends toward over-engineering despite instructions to keep simple; RA converges faster")
para(doc_en, "4. Agent + BO synergy: agent provides direction (structural decisions), BO provides coverage (continuous details)")
para(doc_en, "5. RA made intuitive changes: attention-pooling, Huber loss, orthogonal initialization")
para(doc_en, "6. UA discovered Mixup regularization (2018 ICLR paper) - something the author was not originally familiar with")

h2(doc_en, "Limitations & Future Directions")
para(doc_en, "- Small sample size (3 replicates), modest budget (100 iterations)")
para(doc_en, "- Evaluation quality: single scalar (MAE) may compress too much info")
para(doc_en, "- Context window limits at scale (1000+ iterations)")
para(doc_en, "- Human oversight remains necessary")

h2(doc_en, "Code")
para(doc_en, "https://github.com/DimensionCap/autoresearch_thermo")

doc_en.save('E:/Harry AI 学习/99-论文与阅读/Dimension-TheMeltingPoint/The-Melting-Point-原文版.docx')
print("Melting Point 原文版 done")

# ===================== 2. 中文版 =====================
doc_cn = Document()
style(doc_cn)
h1(doc_cn, "熔点 —— 用 AutoResearch 预测蛋白质热稳定性")
para(doc_cn, "作者: Frank Gao (Dimension Research)")
para(doc_cn, "来源: https://research.dimensioncap.com/p/the-melting-point")
para(doc_cn, "发表时间: 2026-03-23")
para(doc_cn, "")

h2(doc_cn, "摘要")
para(doc_cn, "我们将 autoresearch 框架应用于蛋白质热稳定性预测，给 AI Agent 两种不同自由度的设置，并与贝叶斯优化进行对比。我们的结论是：直觉与穷举可以协同工作。")

h2(doc_cn, "背景：为什么热稳定性很重要")
para(doc_cn, "每个蛋白质都有一个熔解温度（Tm）。加热超过这个阈值，蛋白质结构就会解体。如果治疗性抗体在 55 摄氏度以下就变性，就无法配制成稳定的药物产品。热稳定性与药物可开发性密切相关：溶解度、抗聚集性、保质期和可制造性。")

para(doc_cn, "相比结合亲和力预测，热稳定性预测受到的关注较少。从还原论角度看，Tm 可以可靠测量，且输出是一个正标量——这使它成为生物学 autoresearch 的理想首个项目。")

h2(doc_cn, "数据集：Meltome Atlas + TemBERTureDB")
para(doc_cn, "- Meltome Atlas：涵盖 13 个物种、超过 13000 种蛋白质的熔解曲线")
para(doc_cn, "- TemBERTureDB：清洗和策划后的版本，平衡了嗜热和非嗜热序列")
para(doc_cn, "- 训练/验证/测试划分：80/10/10，通过聚类防止数据泄漏")
para(doc_cn, "- 任务：从氨基酸序列预测 Tm，主要用摄氏度平均绝对误差（MAE）评估")

h2(doc_cn, "基线模型")
para(doc_cn, "参考模型：TemBERTure_Tm，测试集 MAE 6.42 摄氏度，使用 protBERT-BFD 骨干网络 + Houlsby 瓶颈适配器（在 4.2 亿参数冻结骨干网络上还有 2500 万可训练参数）。")

para(doc_cn, "我们的模型：冻结的 ESM-2 150M 嵌入 + 轻量级自注意力适配器 + MLP 回归头（约 170 万参数，小 15 倍）。")

h2(doc_cn, "五种优化方法")
para(doc_cn, "1. 贝叶斯优化（BO）：非 Agent 基线，通过统计模型搜索约 25 个超参数")
para(doc_cn, "2. 受限 Agent（RA）：与 BO 相同的 25 个旋钮，但由 LLM 控制，每次试验提议 1-3 个更改，跟踪历史")
para(doc_cn, "3. 无约束 Agent（UA）：无任何约束，可以重写整个训练脚本，发明新架构")
para(doc_cn, "4. RA + BO：Agent 确定方向，BO 精细搜索连续子空间")
para(doc_cn, "5. UA + BO：将相同的综合理念应用于 UA 更开放式的提议")

h2(doc_cn, "关键发现")
para(doc_cn, "1. 所有方法用小 15 倍的模型击败了基线")
para(doc_cn, "2. UA 达到最低验证集 MAE；RA+BO 达到最低测试集 MAE（泛化最好）")
para(doc_cn, "3. 自由的悖论：UA 趋向过度工程化，尽管有保持简单的指示；RA 反而收敛更快")
para(doc_cn, "4. Agent + BO 协同：Agent 提供方向（结构决策），BO 提供覆盖（连续细节）")
para(doc_cn, "5. RA 做出了直觉性修改：注意力池化、Huber 损失函数、正交初始化")
para(doc_cn, "6. UA 发现了 Mixup 正则化（2018 年 ICLR 论文），这是作者最初不熟悉的技术")

h2(doc_cn, "局限性与未来方向")
para(doc_cn, "- 样本量小（3 个重复），预算适中（100 次迭代）")
para(doc_cn, "- 评估质量：单一标量（MAE）可能压缩了太多信息")
para(doc_cn, "- 上下文窗口限制：扩展到 1000 次以上迭代时会遇到瓶颈")
para(doc_cn, "- 人类监督仍然必要")

h2(doc_cn, "代码开源")
para(doc_cn, "https://github.com/DimensionCap/autoresearch_thermo")

doc_cn.save('E:/Harry AI 学习/99-论文与阅读/Dimension-TheMeltingPoint/The-Melting-Point-中文版.docx')
print("Melting Point 中文版 done")
