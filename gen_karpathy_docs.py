# -*- coding: utf-8 -*-
"""生成 Karpathy AutoResearch 原文版 + 中文版 docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def style(doc):
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/mathControls}mx', 'Microsoft YaHei')

def h1(doc, t):
    p = doc.add_heading(level=1)
    r = p.add_run(t)
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0,0,128)
    return p

def h2(doc, t):
    p = doc.add_heading(level=2)
    r = p.add_run(t)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0,80,0)
    return p

def para(doc, t, sz=11):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(sz)
    return p

# ===================== 1. 原文版 =====================
doc_en = Document()
style(doc_en)
h1(doc_en, "AutoResearch by Andrej Karpathy")
para(doc_en, "Source: https://github.com/karpathy/autoresearch")
para(doc_en, "")

h2(doc_en, "What is AutoResearch?")
para(doc_en, 'AutoResearch is a project where an AI agent is given a small LLM training codebase (train.py) and a natural language specification (program.md). The agent then iteratively modifies the code, trains the model, evaluates the results, and retains improvements — completely autonomously, running experiments around the clock.')

h2(doc_en, "Core Components")
para(doc_en, "1. train.py — The training script that the agent modifies")
para(doc_en, "2. prepare.py — Data preparation script (agent does NOT modify this)")
para(doc_en, "3. program.md — Natural language instructions for the agent")
para(doc_en, "4. Time budget: 5 minutes per experiment")

h2(doc_en, "The Loop")
para(doc_en, "- Agent reads current train.py")
para(doc_en, "- Agent proposes 1-3 modifications")
para(doc_en, "- Runs training with 5-minute budget")
para(doc_en, "- Evaluates val_bpb (bits per byte)")
para(doc_en, "- If improved: keep changes; if worse: revert")
para(doc_en, "- Repeat for ~100 experiments overnight")

h2(doc_en, "Results")
para(doc_en, "Karpathy ran ~700 experiments over 2 days and achieved ~11% reduction in 'time-to-GPT-2' through stacking ~20 incremental improvements: attention scaler sharpening, missing regularization, optimizer beta fixes, weight decay schedule tuning.")

h2(doc_en, "Small GPU Advice (RTX 2070 class)")
para(doc_en, "- Reduce vocab size (smaller tokenizer)")
para(doc_en, "- Reduce sequence length")
para(doc_en, "- Use TinyStories dataset")
para(doc_en, "- Reduce batch size / gradient accumulation")

doc_en.save('E:/Harry AI 学习/99-论文与阅读/Karpathy-AutoResearch/Karpathy-AutoResearch-原文版.docx')
print("原文版 done")

# ===================== 2. 中文版 =====================
doc_cn = Document()
style(doc_cn)
h1(doc_cn, "AutoResearch — Andrej Karpathy 的自主研究框架")

para(doc_cn, "来源: https://github.com/karpathy/autoresearch")
para(doc_cn, "")

h2(doc_cn, "什么是 AutoResearch？")
para(doc_cn, "AutoResearch 是一个让 AI Agent 自主运行 LLM 训练实验的项目。具体来说，给 Agent 一个小型 LLM 训练代码库（train.py）和一份自然语言规范（program.md），Agent 就会迭代修改代码、训练模型、评估结果、保留改进——完全自主进行，24小时不间断跑实验。")

h2(doc_cn, "核心组成部分")
para(doc_cn, "1. train.py — 训练脚本，Agent 可以修改这个文件")
para(doc_cn, "2. prepare.py — 数据准备脚本（Agent 不能修改这个）")
para(doc_cn, "3. program.md — 给 Agent 的自然语言指令")
para(doc_cn, "4. 时间预算：每次实验固定 5 分钟")

h2(doc_cn, "工作流程")
para(doc_cn, "- Agent 读取当前的 train.py")
para(doc_cn, "- Agent 提议 1-3 个修改")
para(doc_cn, "- 用 5 分钟预算运行训练")
para(doc_cn, "- 评估 val_bpb（每字节比特数）")
para(doc_cn, "- 如果有改进：保留改动；如果变差：回滚")
para(doc_cn, "- 一晚上大约跑 100 次实验")

h2(doc_cn, "实验结果")
para(doc_cn, "Karpathy 在 2 天内跑了约 700 次实验，通过叠加约 20 项增量改进，实现了「达到 GPT-2 水平所需时间」减少约 11%。这些改进包括：锐化注意力缩放器、添加缺失的正则化、修复优化器 beta 值、调整权重衰减计划等。")

h2(doc_cn, "小显存 GPU 的建议（RTX 2070 级别）")
para(doc_cn, "- 减小词表大小（使用更小的 tokenizer）")
para(doc_cn, "- 减小序列长度")
para(doc_cn, "- 使用 TinyStories 数据集")
para(doc_cn, "- 减小 batch size / 使用梯度累积")

doc_cn.save('E:/Harry AI 学习/99-论文与阅读/Karpathy-AutoResearch/Karpathy-AutoResearch-中文版.docx')
print("中文版 done")
