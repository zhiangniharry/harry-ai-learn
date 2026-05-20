# -*- coding: utf-8 -*-
"""
生成课程 1.4 的 Word 讲义
运行: python generate_docx_1.4.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
# 辅助函数
# ============================================================
def set_chinese_font(run, font_name='微软雅黑', font_size=11, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, '微软雅黑', 18, bold=True)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    elif level == 2:
        set_chinese_font(run, '微软雅黑', 14, bold=True)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    else:
        set_chinese_font(run, '微软雅黑', 12, bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11, color=None, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_chinese_font(run, '微软雅黑', size, bold, italic)
    if color:
        run.font.color.rgb = color
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    set_chinese_font(run, 'Consolas', 10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('提示: ' + text)
    set_chinese_font(run, '微软雅黑', 10.5, italic=True)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_table_3col(doc, headers, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_chinese_font(r, '微软雅黑', 11, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_chinese_font(r, '微软雅黑', 10)
    return table

# ============================================================
# 创建文档
# ============================================================
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# ============================================================
# 标题页
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('课程 1.4: MNIST 手写数字分类', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_chinese_font(run, '微软雅黑', 26, bold=True)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CNN 卷积神经网络 / 图像分类 / 训练实战')
set_chinese_font(run, '微软雅黑', 14)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Harry AI 学习课程 · 计算机视觉篇 · 课程 1.4\n2026-05-15')
set_chinese_font(run, '微软雅黑', 11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_heading_zh(doc, '目录', level=1)
toc_items = [
    '一、本节目标：让计算机"看懂"手写数字',
    '二、数据集介绍：MNIST',
    '三、CNN 卷积神经网络是什么？',
    '四、CNN 的核心组成：卷积层、池化层、全连接层',
    '五、超参数与训练配置',
    '六、完整训练代码解读',
    '七、代码文件说明',
    '八、常见问题与调试',
    '九、本节小结',
    '十、作业与思考题',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

doc.add_page_break()

# ============================================================
# 第一章
# ============================================================
add_heading_zh(doc, '一、本节目标：让计算机"看懂"手写数字', level=1)

add_paragraph_zh(doc,
    '你有没有想过，手机输入法里那个"手写输入"功能，是怎么实现的？'
    '背后的技术，就是图像分类。本节我们要用 PyTorch 从头训练一个模型，'
    '让它能准确识别 0-9 这 10 个手写数字。')
add_paragraph_zh(doc, '学完本节，你将掌握：', bold=True)
goals = [
    '如何使用 torchvision 加载 MNIST 数据集',
    'CNN（卷积神经网络）的基本结构和原理',
    '如何定义一个多层神经网络模型',
    '完整的训练流程（前向传播 → 损失计算 → 反向传播 → 参数更新）',
    '如何在测试集上评估模型',
    '如何保存和加载模型',
]
for g in goals:
    p = doc.add_paragraph(g, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_note(doc, '本节代码基于课程 1.2 的 DataLoader 和课程 1.3 的训练循环，'
         '是前面两节的综合实战。有遗忘可以回去翻一翻。')

# ============================================================
# 第二章
# ============================================================
add_heading_zh(doc, '二、数据集介绍：MNIST', level=1)

add_paragraph_zh(doc,
    'MNIST 是深度学习界的"Hello World"。'
    '它包含 70,000 张手写数字图片，每张是 28x28 像素的灰度图。')

add_heading_zh(doc, '2.1 数据集结构', level=2)
add_table_3col(doc,
    ['属性', '训练集', '测试集'],
    [
        ['图片数量', '60,000 张', '10,000 张'],
        ['图片尺寸', '28 x 28 像素', '28 x 28 像素'],
        ['颜色', '灰度（1 通道）', '灰度（1 通道）'],
        ['数字类别', '0, 1, 2, ..., 9', '0, 1, 2, ..., 9'],
    ])

add_heading_zh(doc, '2.2 为什么要用 MNIST？', level=2)
reasons = [
    '简单：28x28 不大不小，适合入门',
    '干净：已经预处理好了，不用操心格式问题',
    '经典：几乎所有深度学习教程都用它，方便对照学习',
    '快速：训练快，CPU 也能几分钟跑完',
    '有意义：能真正体验到模型从"不会"到"会"的全过程',
]
for r in reasons:
    p = doc.add_paragraph(r, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_heading_zh(doc, '2.3 PyTorch 中的 MNIST', level=2)
add_paragraph_zh(doc, '用 torchvision 加载 MNIST 只需要这几行代码：')

add_code_block(doc, '''from torchvision import datasets, transforms

transform = transforms.Compose([
    transforms.ToTensor(),                  # 转成张量 (0-255 → 0-1)
    transforms.Normalize((0.1307,), (0.3081,))  # MNIST 专用均值/标准差
])

train_dataset = datasets.MNIST(
    root='./datasets/MNIST',
    train=True,
    transform=transform,
    download=True   # 第一次自动下载
)

test_dataset = datasets.MNIST(
    root='./datasets/MNIST',
    train=False,
    transform=transform,
    download=True
)''')

add_note(doc, 'Normalize 里的 (0.1307,) 和 (0.3081,) 是 MNIST 训练集的全局均值和标准差。'
         'MNIST 只有灰度一个通道，所以均值和标准差各只有一个值，不是三个。')

# ============================================================
# 第三章
# ============================================================
add_heading_zh(doc, '三、CNN 卷积神经网络是什么？', level=1)

add_paragraph_zh(doc,
    'CNN（Convolutional Neural Network，卷积神经网络）'
    '是计算机视觉领域最核心的模型架构。它模仿人类视觉皮层的工作方式，'
    '能自动从图像中提取特征。')

add_heading_zh(doc, '3.1 图像识别的困难', level=2)
add_paragraph_zh(doc,
    '让计算机识别手写数字很难，因为：\n'
    '- 每个人写的字形状不一样（有人写得方，有人写得圆）\n'
    '- 位置不固定（数字可能偏左、偏右、偏上、偏下）\n'
    '- 大小不固定（有的数字大，有的数字小）\n'
    '- 粗细不固定（笔画有粗有细）')

add_heading_zh(doc, '3.2 传统方法 vs CNN', level=2)
add_table_3col(doc,
    ['对比项', '传统方法', 'CNN（卷积神经网络）'],
    [
        ['特征提取', '人工设计（手工提取HOG、SIFT等）', '自动从数据中学习'],
        ['位置敏感', '需要滑动窗口等技巧', '天然具有平移不变性'],
        ['效果', '在复杂任务上效果差', '效果好，是视觉任务的主流'],
        ['参数效率', '特征少还好，多了计算量爆炸', '权重共享，参数量可控'],
    ])

add_heading_zh(doc, '3.3 CNN 的核心思想', level=2)
add_paragraph_zh(doc,
    'CNN 的核心思想是"局部感受野 + 权重共享"。\n\n'
    '局部感受野：CNN 不一次性看整张图，而是用一个小窗口（如 3x3）逐块扫描。'
    '就像你看一个汉字，是从局部笔画组合来判断的，不是整体模糊一看就知道。\n\n'
    '权重共享：同一个卷积核扫遍整张图，所有位置共用同一套权重。'
    '这大大减少了参数数量，也让模型能识别任意位置的同一特征。')

add_heading_zh(doc, '3.4 类比：识别一只猫', level=2)
add_paragraph_zh(doc,
    'CNN 识别一只猫，就像你这样看：\n'
    '1. 先看局部：有没有尖耳朵？（卷积层找边缘、纹理）\n'
    '2. 组合信息：尖耳朵 + 圆脸 = 猫头？\n'
    '3. 抽象判断：猫头 + 尾巴 + 爪子 = 是一只手')

# ============================================================
# 第四章
# ============================================================
add_heading_zh(doc, '四、CNN 的核心组成：卷积层、池化层、全连接层', level=1)

add_heading_zh(doc, '4.1 卷积层（Conv2d）', level=2)
add_paragraph_zh(doc,
    '卷积层是 CNN 的核心。卷积核（Kernel）是一个小矩阵（如 3x3），'
    '它在图像上滑动，每一步做一次"对应位置相乘再求和"的运算。')

add_table_3col(doc,
    ['参数', '含义', '本课程取值'],
    [
        ['in_channels', '输入通道数', '灰度图 = 1，彩色图 = 3'],
        ['out_channels', '输出通道数（卷积核个数）', '32, 64'],
        ['kernel_size', '卷积核大小', '3x3（最常用）'],
        ['padding', '边缘填充（保持尺寸）', '1（让输入输出尺寸不变）'],
        ['stride', '滑动步长', '1（默认）'],
    ])

add_heading_zh(doc, '4.2 激活函数（ReLU）', level=2)
add_paragraph_zh(doc,
    '激活函数给网络引入非线性。没有它，多层网络退化成一层线性变换，'
    '什么都学不了。ReLU 是最常用的激活函数：')
add_code_block(doc, 'ReLU(x) = max(0, x)  # 小于 0 的变成 0，大于 0 的保持不变')
add_paragraph_zh(doc, '优点：计算极快（就一个比较），梯度消失问题轻。')

add_heading_zh(doc, '4.3 池化层（MaxPool2d）', level=2)
add_paragraph_zh(doc,
    '池化层（Pooling）的作用是缩小图片尺寸，减少计算量，提取主要特征。'
    '最大池化（Max Pooling）在每个小窗口里取最大值：')
add_code_block(doc, '''# kernel_size=2, stride=2 的池化
# 输入: 4x4 矩阵
# 输出: 2x2 矩阵，每块取最大值

[[1, 2, 1, 2],     [[2, 2],
 [3, 4, 3, 4],  →   [4, 4]]
 [1, 2, 1, 2],
 [3, 4, 3, 4]]''')
add_paragraph_zh(doc, '池化后图片尺寸减半：28x28 → 14x14 → 7x7')

add_heading_zh(doc, '4.4 全连接层（Linear）', level=2)
add_paragraph_zh(doc,
    '全连接层（Fully Connected Layer）把前面提取的特征综合起来，做最终分类。'
    '就像厨师把切好的食材（特征）按照食谱组合成一道菜（输出）。')

add_heading_zh(doc, '4.5 整体数据流', level=2)
add_code_block(doc, '''输入图片:  (batch, 1, 28, 28)         # 批次×通道×高×宽
  ↓ Conv1 + ReLU + MaxPool
中间特征: (batch, 32, 14, 14)         # 32 个特征图，每个 14×14
  ↓ Conv2 + ReLU + MaxPool
中间特征: (batch, 64, 7, 7)            # 64 个特征图，每个 7×7
  ↓ Flatten（展平）
一维向量: (batch, 3136)                # 7×7×64 = 3136
  ↓ FC1 + ReLU + Dropout
特征向量: (batch, 128)
  ↓ FC2
输出logit: (batch, 10)                 # 10 个数字的预测分数''')

add_note(doc, 'Dropout 是在训练时随机"关闭"一部分神经元（如 25%），'
         '让模型不依赖某些特定神经元，从而提升泛化能力，防止过拟合。'
         '测试时 Dropout 不起作用。')

add_heading_zh(doc, '4.6 模型参数量估算', level=2)
add_table_3col(doc,
    ['层', '计算方式', '参数量'],
    [
        ['Conv1', '3×3×1×32 + 32', '320'],
        ['Conv2', '3×3×32×64 + 64', '18,496'],
        ['FC1', '3136×128 + 128', '401,536'],
        ['FC2', '128×10 + 10', '1,290'],
        ['总计', '-', '约 42 万参数'],
    ])

add_note(doc, '42 万参数在 CNN 里算很小的。ResNet、VGG 等大模型有几千万到上亿参数。'
         'MNIST 用简单的 CNN 就够了，大模型反而容易过拟合。')

# ============================================================
# 第五章
# ============================================================
add_heading_zh(doc, '五、超参数与训练配置', level=1)

add_heading_zh(doc, '5.1 核心超参数', level=2)
add_table_3col(doc,
    ['超参数', '本课程取值', '说明'],
    [
        ['BATCH_SIZE', '64', '每批 64 张图片；太大显存不够，太小梯度不稳定'],
        ['EPOCHS', '5', '训练 5 轮；MNIST 5 轮通常就能达到 98%+ 准确率'],
        ['LEARNING_RATE', '0.001', 'Adam 的默认学习率，通常不用改'],
        ['DEVICE', 'cuda/cpu', '自动检测，有 GPU 用 GPU，没有用 CPU'],
    ])

add_heading_zh(doc, '5.2 损失函数与优化器', level=2)
add_paragraph_zh(doc,
    '多分类问题用 CrossEntropyLoss（交叉熵损失）：\n'
    '它内部自动做了 Softmax，把网络输出的 10 个分数转成概率分布（和为 1）。\n'
    '优化器用 Adam：自动调节每个参数的学习率，比 SGD 更省心。')

add_code_block(doc, '''criterion = nn.CrossEntropyLoss()          # 多分类损失
optimizer = optim.Adam(model.parameters(), lr=0.001)  # Adam 优化器''')

add_heading_zh(doc, '5.3 训练耗时参考', level=2)
add_table_3col(doc,
    ['硬件', '每轮耗时', '5 轮总耗时'],
    [
        ['RTX 3080 / 3090', '约 3-5 秒', '约 20 秒'],
        ['RTX 3060 / 2060', '约 10-15 秒', '约 60 秒'],
        ['GTX 1080', '约 20-30 秒', '约 2 分钟'],
        ['CPU（i5/i7）', '约 2-5 分钟', '约 10-25 分钟'],
    ])

# ============================================================
# 第六章
# ============================================================
add_heading_zh(doc, '六、完整训练代码解读', level=1)

add_heading_zh(doc, '6.1 超参数设置', level=2)
add_code_block(doc, '''BATCH_SIZE = 64       # 每批处理64张图片
EPOCHS = 5            # 训练5轮
LEARNING_RATE = 0.001 # Adam 默认学习率
DEVICE = torch.device('cuda' if torch.cuda.is_available() else 'cpu')''')

add_heading_zh(doc, '6.2 模型定义', level=2)
add_code_block(doc, '''class SimpleCNN(nn.Module):
    def __init__(self):
        super(SimpleCNN, self).__init__()
        self.conv1 = nn.Conv2d(1, 32, kernel_size=3, padding=1)
        self.conv2 = nn.Conv2d(32, 64, kernel_size=3, padding=1)
        self.fc1   = nn.Linear(7*7*64, 128)
        self.fc2   = nn.Linear(128, 10)
        self.dropout = nn.Dropout(p=0.25)
        self.relu  = nn.ReLU()
        self.maxpool = nn.MaxPool2d(2, 2)

    def forward(self, x):
        x = self.maxpool(self.relu(self.conv1(x)))   # 28→14
        x = self.maxpool(self.relu(self.conv2(x)))   # 14→7
        x = x.view(x.size(0), -1)                   # 展平
        x = self.dropout(self.relu(self.fc1(x)))
        x = self.fc2(x)                             # 输出10个logit
        return x''')

add_note(doc, '注意 fc1 的输入是 7*7*64=3136，这来自 Conv2 输出的尺寸。'
         '7×7 来自 28 经过两次 2×2 池化后：28÷2÷2=7。'
         '如果改变了池化或卷积的参数，这里也要相应调整！')

add_heading_zh(doc, '6.3 训练函数', level=2)
add_code_block(doc, '''def train_epoch(model, train_loader, criterion, optimizer, device):
    model.train()   # 训练模式（启用Dropout）
    running_loss, correct, total = 0.0, 0, 0
    for batch_idx, (data, target) in enumerate(train_loader):
        data, target = data.to(device), target.to(device)

        output = model(data)               # 1. 前向传播
        loss = criterion(output, target)  # 2. 计算损失

        optimizer.zero_grad()             # 3. 清零梯度
        loss.backward()                   # 4. 反向传播
        optimizer.step()                  # 5. 更新参数

        running_loss += loss.item()
        _, predicted = output.max(1)
        total += target.size(0)
        correct += predicted.eq(target).sum().item()

        if (batch_idx + 1) % 100 == 0:
            acc = 100. * correct / total
            print(f"  批次 {batch_idx+1:3d}/{len(train_loader)} | "
                  f"损失 {loss.item():.4f} | 准确率 {acc:.2f}%")

    return running_loss/len(train_loader), 100.*correct/total''')

add_heading_zh(doc, '6.4 测试函数', level=2)
add_code_block(doc, '''def test(model, test_loader, criterion, device):
    model.eval()    # 评估模式（禁用Dropout）
    test_loss, correct, total = 0.0, 0, 0
    with torch.no_grad():   # 不计算梯度，省显存
        for data, target in test_loader:
            data, target = data.to(device), target.to(device)
            output = model(data)
            test_loss += criterion(output, target).item()
            _, predicted = output.max(1)
            total += target.size(0)
            correct += predicted.eq(target).sum().item()

    test_loss = test_loss / len(test_loader)
    test_acc = 100. * correct / total
    return test_loss, test_acc''')

add_heading_zh(doc, '6.5 保存最佳模型', level=2)
add_code_block(doc, '''if test_acc > best_test_acc:
    best_test_acc = test_acc
    torch.save({
        'epoch': epoch,
        'model_state_dict': model.state_dict(),
        'optimizer_state_dict': optimizer.state_dict(),
        'test_acc': test_acc,
    }, 'E:/Harry AI 学习/models/mnist_cnn_best.pth')
    print(f"  ★ 保存最佳模型，准确率 {test_acc:.2f}%")''')

add_note(doc, '推荐保存整个 checkpoint（包含优化器状态），'
         '这样可以从中断的地方继续训练（resume）。'
         '如果只保存 state_dict，加载后需要重新创建优化器。')

# ============================================================
# 第七章
# ============================================================
add_heading_zh(doc, '七、代码文件说明', level=1)
add_table_3col(doc,
    ['文件名', '作用', '说明'],
    [
        ['04-mnist_cnn_train.py', '主训练脚本', '包含数据加载、模型定义、训练、测试、保存完整流程'],
        ['generate_docx_1.4.py', '讲义生成脚本', '运行后生成课程 1.4 讲义 Word 文档'],
        ['mnist_cnn_best.pth', '训练好的模型', '准确率最高的模型权重（保存在 models/ 目录）'],
    ])

add_heading_zh(doc, '运行方式', level=2)
add_code_block(doc, '''# 在项目根目录运行
cd "E:/Harry AI 学习"
python "03-计算机视觉/01-MNIST分类/04-mnist_cnn_train.py"''')

add_note(doc, 'Windows 下 num_workers 建议设为 0，避免多进程问题。'
         'Linux/Mac 可以设为 2 或 4，加速数据加载。')

# ============================================================
# 第八章
# ============================================================
add_heading_zh(doc, '八、常见问题与调试', level=1)
add_table_3col(doc,
    ['问题', '原因', '解决方案'],
    [
        ['Loss 不下降', '学习率太大或太小', '尝试 lr=0.01 或 lr=0.0001'],
        ['准确率停在 10%', '模型没学到东西', '检查数据预处理是否正确，确认 DEVICE 设置'],
        ['显存不够（OOM）', 'BATCH_SIZE 太大', '把 BATCH_SIZE 从 64 降到 32 或 16'],
        ['训练太慢', '用 CPU 而非 GPU', '确认 torch.cuda.is_available()=True'],
        ['模型加载报错', '保存和加载的结构不一致', '确保加载前先创建相同结构的模型'],
        ['RuntimeError: expected...', '维度不匹配', '检查 fc1 的输入维度是否和前面的输出匹配'],
    ])

add_heading_zh(doc, '8.1 快速调试技巧', level=2)
tips = [
    '先把 BATCH_SIZE 设为 1，看输入输出维度是否匹配',
    '第一次跑用 CPU，确认逻辑正确后再切到 GPU',
    '加 print(model) 快速检查模型结构',
    '加 print(x.shape) 查看中间张量形状',
    '损失不降时，把学习率改成 0.01 或 0.0001 重试',
]
for t in tips:
    p = doc.add_paragraph(t, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

# ============================================================
# 第九章
# ============================================================
add_heading_zh(doc, '九、本节小结', level=1)

concepts = [
    ('MNIST', '手写数字数据集', '70,000张 28x28 灰度图，深度学习入门神器'),
    ('CNN', '卷积神经网络', '自动提取图像特征，平移不变性'),
    ('卷积层', 'Conv2d', '用小窗口提取局部特征'),
    ('池化层', 'MaxPool2d', '缩小尺寸，减少计算，提取主干特征'),
    ('Dropout', '随机丢弃', '防止过拟合，提升泛化能力'),
    ('CrossEntropyLoss', '交叉熵损失', '多分类问题的标准损失函数'),
    ('Adam', '自适应优化器', '自动调学习率，最常用的优化器'),
    ('Checkpoint', '检查点保存', '保存模型+优化器状态，支持恢复训练'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['概念', '代码/英文', '说明']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            set_chinese_font(r, '微软雅黑', 11, bold=True)
for row_data in concepts:
    cells = table.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val
        for p in cells[i].paragraphs:
            for r in p.runs:
                set_chinese_font(r, '微软雅黑', 10)

add_paragraph_zh(doc, '\n预期训练结果：', bold=True)
add_paragraph_zh(doc,
    '运行 5 轮后，测试准确率应该在 97%-99% 之间。'
    'CNN 模型在这个任务上非常高效，能轻松达到 98%+ 的准确率。'
    '如果你跑出来准确率很低（<90%），回过头检查数据预处理是否正确。')

# ============================================================
# 第十章
# ============================================================
add_heading_zh(doc, '十、作业与思考题', level=1)

add_heading_zh(doc, '作业', level=2)
assignments = [
    '运行 04-mnist_cnn_train.py，记录每轮的准确率',
    '把 EPOCHS 改成 10，对比准确率变化',
    '把 BATCH_SIZE 改成 32，观察训练速度变化',
    '把学习率从 0.001 改成 0.01，观察损失曲线有什么不同',
    '运行完成后，找到 mnist_cnn_best.pth，查看文件大小',
]
for i, a in enumerate(assignments, 1):
    p = doc.add_paragraph(f'{i}. {a}')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_heading_zh(doc, '思考题', level=2)
questions = [
    '为什么 MNIST 图片是 (1, 28, 28) 而不是 (28, 28, 1)？PyTorch 对通道顺序有什么要求？',
    'CNN 里的 padding=1 有什么作用？如果 padding=0，图片尺寸会怎么变化？',
    'Dropout 为什么能防止过拟合？测试时为什么 Dropout 不起作用？',
    '如果 Conv2 输出的特征图不是 7x7 而是 8x8，fc1 的输入维度应该改成多少？',
    '模型训练到后面，准确率涨得很慢甚至不动了，这说明什么？可以怎么改进？',
]
for q in questions:
    p = doc.add_paragraph(q, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_heading_zh(doc, '下一节预告', level=2)
add_paragraph_zh(doc,
    '课程 2.1：CIFAR-10 彩色图像分类\n'
    'MNIST 只是灰度单通道的小图片。下一节我们升级到 CIFAR-10，'
    '10 个类别（猫、狗、飞机等）的彩色图片，用真实的彩色图像继续训练！'
    '你会学到如何处理 3 通道彩色图片，以及更深层的网络结构。',
    color=RGBColor(0x00, 0x66, 0x99))

# ============================================================
# 保存文档
# ============================================================
output_path = r'E:\Harry AI 学习\03-计算机视觉\01-MNIST分类\课程1.4-MNIST手写数字分类.docx'
doc.save(output_path)
print(f'讲义已生成: {output_path}')
