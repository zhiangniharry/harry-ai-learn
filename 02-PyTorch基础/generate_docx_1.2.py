# -*- coding: utf-8 -*-
"""
生成课程 1.2 的 Word 文档
运行: python generate_docx_1.2.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='微软雅黑', font_size=11, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, '微软雅黑', 18, bold=True)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    elif level == 2:
        set_chinese_font(run, '微软雅黑', 14, bold=True)
    else:
        set_chinese_font(run, '微软雅黑', 12, bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11, color=None):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '微软雅黑', size, bold)
    if color:
        run.font.color.rgb = color
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    set_chinese_font(run, 'Consolas', 10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    return p

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ========== 标题页 ==========
title = doc.add_heading('课程 1.2: PyTorch 数据管道', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_chinese_font(run, '微软雅黑', 24, bold=True)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Dataset / DataLoader / 预处理 / 数据增强')
set_chinese_font(run, '微软雅黑', 14)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Harry AI 学习课程\n2026-05-15')
set_chinese_font(run, '微软雅黑', 11)

doc.add_page_break()

# ========== 目录 ==========
add_heading_zh(doc, '目录', level=1)
toc_items = [
    '一、什么是数据管道？',
    '二、Dataset（数据集）',
    '三、DataLoader（数据加载器）',
    '四、图像数据预处理',
    '五、数据增强（Data Augmentation）',
    '六、torchvision.transforms 实战',
    '七、总结与作业'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

doc.add_page_break()

# ========== 第一章 ==========
add_heading_zh(doc, '一、什么是数据管道？', level=1)

add_paragraph_zh(doc, 
    '想象你要开一家餐厅。你需要：', size=11)

add_paragraph_zh(doc, 
    '1. 食材（数据）：从市场买来各种原材料\n'
    '2. 厨师切配（预处理）：把食材洗干净、切成统一大小\n'
    '3. 摆盘（DataLoader）：把切好的食材按每桌的份量分好\n'
    '4. 上菜（训练）：一道一道端给客人（模型）',
    size=11)

add_paragraph_zh(doc, 
    '在机器学习中，这个流程就叫"数据管道"（Data Pipeline）。'
    'PyTorch 用两个核心工具来实现：Dataset 和 DataLoader。',
    bold=True, size=11)

add_paragraph_zh(doc, 
    '为什么需要数据管道？\n'
    '- 数据可能很大（几百万张图片），不可能一次性全部加载到内存\n'
    '- 数据需要预处理（调整大小、归一化）\n'
    '- 训练时需要随机打乱顺序，防止模型记住顺序\n'
    '- 需要分批处理（batch），GPU 内存有限',
    size=11)

# ========== 第二章 ==========
add_heading_zh(doc, '二、Dataset（数据集）', level=1)

add_paragraph_zh(doc, 
    'Dataset 就像一本食谱书，告诉 PyTorch 三件事：',
    bold=True, size=11)

add_paragraph_zh(doc, 
    '1. 你一共有多少数据？ → __len__()\n'
    '2. 怎么拿到第 i 个数据？ → __getitem__(i)\n'
    '3. 每个数据长什么样？ → 返回什么',
    size=11)

add_heading_zh(doc, '2.1 类比：餐厅菜单', level=2)
add_paragraph_zh(doc, 
    '去餐厅吃饭时，菜单（Dataset）告诉你：\n'
    '- 一共有 50 道菜（__len__ 返回 50）\n'
    '- 第 3 道菜是宫保鸡丁（__getitem__(2) 返回宫保鸡丁的信息）\n'
    '- 每道菜有名字、价格、图片（返回的是一个结构化的数据）',
    size=11)

add_heading_zh(doc, '2.2 代码实现', level=2)
add_paragraph_zh(doc, '最简单的 Dataset 示例：数字预测', size=11)

code1 = """class SimpleNumberDataset(Dataset):
    def __init__(self, start=0, end=100):
        # 初始化：准备数据
        self.x = torch.arange(start, end, dtype=torch.float32)
        self.y = 2 * self.x + 1  # 标签：y = 2x + 1
    
    def __len__(self):
        # 返回总数据量
        return len(self.x)
    
    def __getitem__(self, idx):
        # 返回第 idx 个数据
        return self.x[idx], self.y[idx]"""
add_code_block(doc, code1)

add_paragraph_zh(doc, 
    '关键点：\n'
    '- __init__：构造函数，准备数据。就像厨师早上进货、备料\n'
    '- __len__：返回数据集大小。就像菜单上写着"本店共 50 道菜"\n'
    '- __getitem__：按索引取数据。就像客人说"我要第 3 道菜"',
    size=11)

add_heading_zh(doc, '2.3 运行结果解读', level=2)
add_paragraph_zh(doc, 
    '当我们创建 dataset = SimpleNumberDataset(start=0, end=10) 时：\n'
    '- 生成了 10 个数据点（x=0,1,2,...,9）\n'
    '- 对应的标签是 y=1,3,5,...,19\n'
    '- dataset[0] 返回 (0.0, 1.0)，即 x=0, y=2*0+1=1',
    size=11)

# ========== 第三章 ==========
add_heading_zh(doc, '三、DataLoader（数据加载器）', level=1)

add_paragraph_zh(doc, 
    'DataLoader 就像一个智能服务员，它的工作是把数据分批、打乱、高效地送给模型。',
    bold=True, size=11)

add_heading_zh(doc, '3.1 类比：餐厅服务员', level=2)
add_paragraph_zh(doc, 
    '假设有 10 桌客人（10 个数据），服务员（DataLoader）的工作：\n'
    '- 分批：每批送 3 桌的菜（batch_size=3）\n'
    '- 打乱：每天上菜的顺序不一样（shuffle=True）\n'
    '- 多线程：多个服务员同时端菜（num_workers>0）',
    size=11)

add_heading_zh(doc, '3.2 关键参数', level=2)

params = [
    ('dataset', '要加载的数据集', '告诉服务员从哪个厨房取菜'),
    ('batch_size', '每批多少数据', '一次端几桌的菜'),
    ('shuffle', '是否打乱顺序', '训练时打乱，测试时不打乱'),
    ('num_workers', '加载数据的进程数', '几个服务员同时工作'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '参数'
hdr_cells[1].text = '含义'
hdr_cells[2].text = '类比'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for param, meaning, analogy in params:
    row_cells = table.add_row().cells
    row_cells[0].text = param
    row_cells[1].text = meaning
    row_cells[2].text = analogy
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_heading_zh(doc, '3.3 代码示例', level=2)
code2 = """dataloader = DataLoader(
    dataset=dataset,      # 数据集
    batch_size=3,         # 每批 3 个
    shuffle=True,         # 打乱顺序
    num_workers=0         # Windows 建议用 0
)

# 遍历数据
for batch_idx, (batch_x, batch_y) in enumerate(dataloader):
    print(f"批次 {batch_idx}: x={batch_x}, y={batch_y}")"""
add_code_block(doc, code2)

add_paragraph_zh(doc, 
    '输出解读：\n'
    '- Epoch 1 的顺序可能是 [7,8,4], [5,1,6], [0,2,3], [9]\n'
    '- Epoch 2 的顺序变了，比如 [9,3,2], [1,7,4], [0,6,5], [8]\n'
    '- 这就是 shuffle=True 的效果：每次 epoch 顺序都不同',
    size=11)

# ========== 第四章 ==========
add_heading_zh(doc, '四、图像数据预处理', level=1)

add_paragraph_zh(doc, 
    '原始图像数据有各种问题，必须预处理才能喂给模型。',
    bold=True, size=11)

add_heading_zh(doc, '4.1 原始数据的问题', level=2)
problems = [
    '尺寸不一样：有的 1000x800，有的 600x400',
    '像素值范围不同：有的是 0-255，有的是 0-1',
    '格式不同：JPG、PNG、BMP',
    '通道顺序不同：RGB vs BGR',
]
for problem in problems:
    p = doc.add_paragraph(problem, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_heading_zh(doc, '4.2 预处理步骤', level=2)
add_paragraph_zh(doc, 
    '1. Resize（调整大小）：把所有图片变成统一尺寸\n'
    '   例如：224x224（这是 ResNet 等模型的标准输入尺寸）\n\n'
    '2. Normalize（归一化）：把像素值缩放到标准范围\n'
    '   例如：(pixel - mean) / std，让数据分布更稳定\n\n'
    '3. ToTensor（转张量）：把 PIL Image 或 numpy 转成 PyTorch Tensor\n'
    '   同时会把值从 [0, 255] 缩放到 [0, 1]\n\n'
    '4. 调整维度顺序：从 (H, W, C) 变成 (C, H, W)\n'
    '   PyTorch 的卷积层要求通道在前',
    size=11)

add_heading_zh(doc, '4.3 为什么用 ImageNet 的均值和标准差？', level=2)
add_paragraph_zh(doc, 
    'ImageNet 是一个包含 120 万张图片、1000 个类别的巨型数据集。'
    '绝大多数预训练模型（如 ResNet、VGG）都是在 ImageNet 上训练的。\n\n'
    '使用 ImageNet 的统计值进行归一化：\n'
    '- mean=[0.485, 0.456, 0.406]\n'
    '- std=[0.229, 0.224, 0.225]\n\n'
    '好处：让你的数据和预训练模型的训练数据分布一致，'
    '这样迁移学习（Transfer Learning）时效果更好。'
    '就像你去国外餐厅，用当地的礼仪会更受欢迎。',
    size=11)

# ========== 第五章 ==========
add_heading_zh(doc, '五、数据增强（Data Augmentation）', level=1)

add_paragraph_zh(doc, 
    '数据增强 = 用现有数据生成"新"数据。'
    '就像给同一个东西拍很多不同角度的照片。',
    bold=True, size=11)

add_heading_zh(doc, '5.1 为什么要做数据增强？', level=2)
benefits = [
    '数据量变多：不用花钱收集新数据',
    '模型更鲁棒：见过各种变形，不容易被欺骗',
    '防止过拟合：不会只记住训练数据的特定样子',
]
for benefit in benefits:
    p = doc.add_paragraph(benefit, style='List Bullet')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

add_heading_zh(doc, '5.2 常见的增强方法', level=2)

aug_methods = [
    ('RandomCrop', '随机裁剪', '从大图切出一小块'),
    ('RandomHorizontalFlip', '随机水平翻转', '50% 概率左右镜像'),
    ('RandomRotation', '随机旋转', '转一个角度（如 ±10 度）'),
    ('ColorJitter', '颜色抖动', '改变亮度、对比度、饱和度'),
    ('Normalize', '标准化', '减均值除标准差'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '方法名'
hdr_cells[1].text = '中文名'
hdr_cells[2].text = '效果'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for method, cn_name, effect in aug_methods:
    row_cells = table2.add_row().cells
    row_cells[0].text = method
    row_cells[1].text = cn_name
    row_cells[2].text = effect
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_heading_zh(doc, '5.3 重要原则', level=2)
add_paragraph_zh(doc, 
    '训练时做增强，测试时不做！\n'
    '- 训练：shuffle=True，用各种增强，让模型见多识广\n'
    '- 测试：shuffle=False，不做增强，保证结果可复现',
    size=11, color=RGBColor(0xC0, 0x00, 0x00))

# ========== 第六章 ==========
add_heading_zh(doc, '六、torchvision.transforms 实战', level=1)

add_paragraph_zh(doc, 
    'torchvision.transforms 是 PyTorch 官方提供的预处理工具，'
    '把各种操作封装成可以链式调用的模块。',
    size=11)

code3 = """from torchvision import transforms

transform = transforms.Compose([
    transforms.ToPILImage(),            # 转成 PIL 格式
    transforms.Resize((224, 224)),      # 调整大小
    transforms.RandomHorizontalFlip(),  # 随机翻转
    transforms.RandomRotation(10),      # 随机旋转
    transforms.ToTensor(),              # 转张量
    transforms.Normalize(               # 标准化
        mean=[0.485, 0.456, 0.406],
        std=[0.229, 0.224, 0.225]
    )
])"""
add_code_block(doc, code3)

add_paragraph_zh(doc, 
    '使用方式：\n'
    'dataset = MyDataset(transform=transform)\n'
    'dataloader = DataLoader(dataset, batch_size=32, shuffle=True)',
    size=11)

# ========== 第七章 ==========
add_heading_zh(doc, '七、总结与作业', level=1)

add_heading_zh(doc, '7.1 核心概念总结', level=2)
concepts = [
    ('Dataset', '菜单', '告诉 PyTorch 数据在哪里，长什么样'),
    ('DataLoader', '服务员', '分批、打乱、高效加载数据'),
    ('预处理', '厨师切菜', '让数据格式统一'),
    ('数据增强', '换角度拍照', '用现有数据生成新数据'),
    ('Compose', '流水线', '把多个操作串起来'),
]

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = '概念'
hdr_cells[1].text = '类比'
hdr_cells[2].text = '作用'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for concept, analogy, role in concepts:
    row_cells = table3.add_row().cells
    row_cells[0].text = concept
    row_cells[1].text = analogy
    row_cells[2].text = role
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_heading_zh(doc, '7.2 作业', level=2)
add_paragraph_zh(doc, 
    '1. 修改代码中的 batch_size，观察输出变化（试试 1, 5, 10）\n'
    '2. 把 shuffle=False，观察两轮 epoch 的顺序是否相同\n'
    '3. 尝试添加新的数据增强方法（如 transforms.RandomCrop）\n'
    '4. 思考：为什么 num_workers 在 Windows 上建议设为 0？',
    size=11)

add_heading_zh(doc, '7.3 下一课预告', level=2)
add_paragraph_zh(doc, 
    '课程 1.3：训练循环\n'
    '我们将用今天学的 DataLoader，写一个完整的训练脚本！\n'
    '包括：前向传播、计算损失、反向传播、更新权重、保存模型。',
    size=11, color=RGBColor(0x00, 0x66, 0x99))

# 保存文档
output_path = r'E:\Harry AI 学习\02-PyTorch基础\课程1.2-数据管道详解.docx'
doc.save(output_path)
print(f'文档已保存: {output_path}')
