# -*- coding: utf-8 -*-
"""
生成课程 1.3 的 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='微软雅黑', font_size=11, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, '微软雅黑', 18, bold=True)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    elif level == 2:
        set_chinese_font(run, '微软雅黑', 14, bold=True)
    else:
        set_chinese_font(run, '微软雅黑', 12, bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '微软雅黑', size, bold)
    if color:
        run.font.color.rgb = color
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    set_chinese_font(run, 'Consolas', 10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    return p

# 创建文档
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题页
title = doc.add_heading('课程 1.3: PyTorch 训练循环', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_chinese_font(run, '微软雅黑', 24, bold=True)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('前向传播 / 损失计算 / 反向传播 / 优化器 / 保存模型')
set_chinese_font(run, '微软雅黑', 14)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Harry AI 学习课程\n2026-05-15')
set_chinese_font(run, '微软雅黑', 11)

doc.add_page_break()

# 目录
add_heading_zh(doc, '目录', level=1)
toc_items = [
    '一、训练循环的 5 个步骤',
    '二、定义模型',
    '三、损失函数',
    '四、优化器',
    '五、完整训练代码',
    '六、保存和加载模型',
    '七、总结'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_chinese_font(run, '微软雅黑', 11)

doc.add_page_break()

# 第一章
add_heading_zh(doc, '一、训练循环的 5 个步骤', level=1)

add_paragraph_zh(doc, 
    '训练一个神经网络，本质上是一个循环。每一轮循环（epoch），'
    '模型都会看一遍所有数据，不断调整自己的参数。',
    size=11)

add_paragraph_zh(doc, '5 个步骤:', bold=True, size=11)

steps = [
    ('1. 前向传播 (Forward)', '输入数据 → 模型计算 → 得到预测结果'),
    ('2. 计算损失 (Loss)', '比较预测值和真实值，算出差多少'),
    ('3. 反向传播 (Backward)', '从损失往回推，计算每个参数的梯度'),
    ('4. 更新权重 (Step)', '用优化器根据梯度调整参数'),
    ('5. 清零梯度 (Zero Grad)', '防止梯度累积，准备下一轮'),
]

for step_name, step_desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(step_name + '\n')
    set_chinese_font(run, '微软雅黑', 11, bold=True)
    run = p.add_run('   ' + step_desc)
    set_chinese_font(run, '微软雅黑', 11)

add_paragraph_zh(doc, 
    '\n类比：学生学习和老师批改作业',
    bold=True, size=11)

add_paragraph_zh(doc, 
    '1. 前向传播 = 学生做题，写出答案\n'
    '2. 计算损失 = 老师批改，算错了多少分\n'
    '3. 反向传播 = 老师分析，哪道题错最多，哪部分知识最薄弱\n'
    '4. 更新权重 = 学生根据反馈，调整学习方法\n'
    '5. 清零梯度 = 学生把上次的错题本收起来，准备做新题',
    size=11)

# 第二章
add_heading_zh(doc, '二、定义模型', level=1)

add_paragraph_zh(doc, 
    '模型就是一个数学函数 y = f(x)。'
    '我们的目标是找到正确的 f，让预测尽量准确。',
    size=11)

add_heading_zh(doc, '2.1 线性层', level=2)
add_paragraph_zh(doc, 
    '最简单的神经网络单元是线性层（Linear Layer），公式是：\n'
    'y = weight * x + bias\n\n'
    '- weight（权重）：决定输入 x 对输出 y 的影响程度\n'
    '- bias（偏置）：决定当 x=0 时，y 的基准值\n\n'
    '对于 y = 2x + 1 这个问题：\n'
    '- 正确的 weight 应该是 2\n'
    '- 正确的 bias 应该是 1\n'
    '- 模型一开始不知道，通过训练自己找出来',
    size=11)

add_code_block(doc, '''class SimpleModel(nn.Module):
    def __init__(self):
        super(SimpleModel, self).__init__()
        # 定义一个线性层: 输入1维, 输出1维
        self.linear = nn.Linear(in_features=1, out_features=1)
    
    def forward(self, x):
        return self.linear(x)''')

add_paragraph_zh(doc, 
    'nn.Module 是 PyTorch 所有神经网络模型的基类。'
    '任何模型都必须继承它，并实现 __init__ 和 forward 方法。',
    size=11)

# 第三章
add_heading_zh(doc, '三、损失函数', level=1)

add_paragraph_zh(doc, 
    '损失函数衡量预测值和真实值之间的差距。'
    '差距越大，损失越大，模型越差。',
    size=11)

add_heading_zh(doc, '3.1 常用损失函数', level=2)

loss_funcs = [
    ('MSELoss', '均方误差', '回归问题（预测连续值，如房价、温度）'),
    ('CrossEntropyLoss', '交叉熵', '分类问题（预测类别，如猫/狗/鸟）'),
    ('L1Loss', '绝对误差', '回归问题，对异常值不敏感'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '损失函数'
hdr_cells[1].text = '中文名'
hdr_cells[2].text = '适用场景'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for name, cn, scene in loss_funcs:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = cn
    row_cells[2].text = scene
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_paragraph_zh(doc, 
    '\n我们的任务是预测 y = 2x + 1，y 是连续值，所以用 MSELoss。',
    size=11)

add_code_block(doc, '''criterion = nn.MSELoss()

# 使用
loss = criterion(predictions, true_values)''')

# 第四章
add_heading_zh(doc, '四、优化器', level=1)

add_paragraph_zh(doc, 
    '优化器根据损失的梯度，更新模型的参数。'
    '类比：爬山时，梯度告诉你哪个方向最陡，'
    '优化器就是沿着最陡的反方向走一步。',
    size=11)

add_heading_zh(doc, '4.1 常用优化器', level=2)

optimizers = [
    ('SGD', '随机梯度下降', '最基础，需要手动调学习率'),
    ('Adam', '自适应矩估计', '最常用，自动调整学习率，效果通常最好'),
    ('RMSprop', '均方根传播', '适合循环神经网络 RNN'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '优化器'
hdr_cells[1].text = '中文名'
hdr_cells[2].text = '特点'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for name, cn, feature in optimizers:
    row_cells = table2.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = cn
    row_cells[2].text = feature
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_paragraph_zh(doc, 
    '\n学习率（lr）是优化器最重要的参数：\n'
    '- 太大：步子迈太大，可能跳过最优解，甚至发散\n'
    '- 太小：步子迈太小，收敛太慢，训练时间太长\n'
    '- 合适：通常从 0.01 或 0.001 开始尝试',
    size=11)

add_code_block(doc, '''optimizer = optim.Adam(model.parameters(), lr=0.01)''')

# 第五章
add_heading_zh(doc, '五、完整训练代码', level=1)

add_paragraph_zh(doc, '训练循环的核心代码:', size=11)

add_code_block(doc, '''for epoch in range(100):  # 训练 100 轮
    model.train()  # 设置训练模式
    
    for batch_x, batch_y in train_loader:
        # 1. 前向传播
        predictions = model(batch_x)
        
        # 2. 计算损失
        loss = criterion(predictions, batch_y)
        
        # 3. 反向传播
        optimizer.zero_grad()  # 清零旧梯度
        loss.backward()        # 计算新梯度
        
        # 4. 更新权重
        optimizer.step()
    
    # 测试（不计算梯度）
    model.eval()
    with torch.no_grad():
        for batch_x, batch_y in test_loader:
            predictions = model(batch_x)
            test_loss = criterion(predictions, batch_y)''')

add_paragraph_zh(doc, 
    '\n关键细节：\n'
    '- model.train()：启用 dropout、batch normalization 等训练专用层\n'
    '- model.eval()：关闭这些层，保证测试/预测结果稳定\n'
    '- torch.no_grad()：测试时不计算梯度，节省内存和计算\n'
    '- optimizer.zero_grad()：必须每轮清零，否则梯度会累积',
    size=11)

# 第六章
add_heading_zh(doc, '六、保存和加载模型', level=1)

add_paragraph_zh(doc, 
    '训练一次可能要几小时甚至几天，保存后可以直接加载使用。',
    size=11)

add_heading_zh(doc, '6.1 保存模型', level=2)
add_code_block(doc, '''# 保存模型参数
torch.save(model.state_dict(), 'model.pth')

# state_dict() 只保存参数（weight、bias等）
# 不保存模型结构，所以加载时需要先创建模型对象''')

add_heading_zh(doc, '6.2 加载模型', level=2)
add_code_block(doc, '''# 1. 创建模型（结构要和保存时一样）
model = SimpleModel()

# 2. 加载参数
model.load_state_dict(torch.load('model.pth'))

# 3. 设置为评估模式
model.eval()''')

add_paragraph_zh(doc, 
    '\n注意：保存和加载时，模型结构必须一致。'
    '如果模型结构变了（比如加了新层），旧的参数就加载不进去了。',
    size=11)

# 第七章
add_heading_zh(doc, '七、总结', level=1)

add_heading_zh(doc, '7.1 核心概念', level=2)

concepts = [
    ('前向传播', '学生做题', '输入 → 模型 → 预测'),
    ('损失函数', '老师批改', '预测 vs 真实，算差距'),
    ('反向传播', '老师分析', '从损失往回算梯度'),
    ('优化器', '学生调整', '根据梯度更新参数'),
    ('训练循环', '日复一日的练习', '重复 1-4 步，直到收敛'),
]

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = '概念'
hdr_cells[1].text = '类比'
hdr_cells[2].text = '作用'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, '微软雅黑', 11, bold=True)

for concept, analogy, role in concepts:
    row_cells = table3.add_row().cells
    row_cells[0].text = concept
    row_cells[1].text = analogy
    row_cells[2].text = role
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '微软雅黑', 10)

add_heading_zh(doc, '7.2 训练效果', level=2)
add_paragraph_zh(doc, 
    '运行代码后，你应该看到：\n'
    '- 训练损失逐渐下降（从几千降到接近 0）\n'
    '- weight 接近 2.0，bias 接近 1.0\n'
    '- 预测值和真实值误差很小（小于 1）',
    size=11)

add_heading_zh(doc, '7.3 下一课预告', level=2)
add_paragraph_zh(doc, 
    '课程 1.4：MNIST 手写数字分类\n'
    '用今天学的训练循环，训练一个真正的神经网络！\n'
    '识别 0-9 的手写数字，CNN 卷积神经网络入门。',
    size=11, color=RGBColor(0x00, 0x66, 0x99))

# 保存
output_path = r'E:\Harry AI 学习\02-PyTorch基础\课程1.3-训练循环详解.docx'
doc.save(output_path)
print(f'文档已保存: {output_path}')
